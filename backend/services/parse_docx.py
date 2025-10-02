from __future__ import annotations

from pathlib import Path

from ..models import ParsedObject

try:  # pragma: no cover - optional dependency
    from docx import Document
except Exception:  # pragma: no cover - dependency missing
    Document = None  # type: ignore


def parse_docx(file_path: str) -> list[ParsedObject]:
    """Parse DOCX files into ParsedObject instances."""

    if Document is None:
        return []

    file_id = Path(file_path).resolve().parent.parent.name
    document = Document(file_path)
    objects: list[ParsedObject] = []
    order_index = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        objects.append(
            ParsedObject(
                object_id=f"{file_id}-docx-text-{order_index:06d}",
                file_id=file_id,
                kind="text",
                text=text,
                page_index=None,
                bbox=None,
                order_index=order_index,
                metadata={"engine": "docx", "source": "python-docx"},
            )
        )
        order_index += 1

    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_text = "\n".join("\t".join(cell for cell in row) for row in rows)
        objects.append(
            ParsedObject(
                object_id=f"{file_id}-docx-table-{table_index:06d}",
                file_id=file_id,
                kind="table",
                text=table_text or None,
                page_index=None,
                bbox=None,
                order_index=order_index,
                metadata={"engine": "docx", "source": "python-docx"},
            )
        )
        order_index += 1

    return objects
