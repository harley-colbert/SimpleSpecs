"""DOCX parsing utilities."""
from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document


def parse_docx(path: Path) -> list[dict[str, Any]]:
    """Parse a DOCX document into normalized objects."""

    document = Document(path)
    objects: list[dict[str, Any]] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        objects.append(
            {
                "line_id": str(uuid4()),
                "type": "text",
                "page": None,
                "bbox": None,
                "content": text,
                "meta": {
                    "style": paragraph.style.name if paragraph.style else None,
                },
            }
        )

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        content_rows = [", ".join(filter(None, row)) for row in rows if any(cell for cell in row)]
        if not content_rows:
            continue
        objects.append(
            {
                "line_id": str(uuid4()),
                "type": "table",
                "page": None,
                "bbox": None,
                "content": "\n".join(content_rows),
                "meta": {"rows": rows},
            }
        )

    return objects
