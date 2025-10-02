from __future__ import annotations

import importlib.util
from io import BytesIO
from typing import Any

import pytest
from fastapi.testclient import TestClient

from backend.main import create_app


client = TestClient(create_app())
DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None


def _assert_objects_shape(objects: list[dict[str, Any]]) -> None:
    assert isinstance(objects, list) and objects, "Expected parsed objects"
    required_keys = {"object_id", "file_id", "kind", "order_index", "metadata"}
    for obj in objects:
        assert required_keys.issubset(obj.keys())
        assert isinstance(obj["metadata"], dict)
    orders = [obj["order_index"] for obj in objects]
    assert orders == sorted(orders)
    assert len(set(orders)) == len(orders)
    ordering_keys = [((obj.get("page_index") or 0), obj["order_index"]) for obj in objects]
    assert ordering_keys == sorted(ordering_keys)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_parse_docx_golden() -> None:
    from docx import Document  # type: ignore

    document = Document()
    document.add_paragraph("Hello")
    document.add_paragraph("World")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    files = {
        "file": (
            "demo.docx",
            buffer,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    response = client.post("/ingest", files=files)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["object_count"] >= 2

    parsed = client.get(f"/parsed/{payload['file_id']}")
    assert parsed.status_code == 200, parsed.text
    objects = parsed.json()
    _assert_objects_shape(objects)
    texts = [obj.get("text") for obj in objects if obj["kind"] == "text" and obj.get("text")]
    assert any(text and ("Hello" in text or "World" in text) for text in texts)
