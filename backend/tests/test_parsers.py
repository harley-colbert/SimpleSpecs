"""Unit tests for document parsers."""
from __future__ import annotations

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from backend.services.parsing import parse_document


def test_parse_txt(tmp_path: Path) -> None:
    sample = tmp_path / "sample.txt"
    sample.write_text("Line one\nLine two\n\nLine three\n", encoding="utf-8")

    objects = parse_document(sample)

    assert len(objects) == 3
    assert all(obj["type"] == "text" for obj in objects)
    assert {obj["content"] for obj in objects} == {"Line one", "Line two", "Line three"}


def test_parse_docx(tmp_path: Path) -> None:
    from docx import Document

    sample = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Introduction")
    document.add_paragraph("Details paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Item"
    table.cell(1, 1).text = "42"
    document.save(sample)

    objects = parse_document(sample)

    texts = [obj for obj in objects if obj["type"] == "text"]
    tables = [obj for obj in objects if obj["type"] == "table"]

    assert any(obj["content"] == "Introduction" for obj in texts)
    assert any("Header" in obj["content"] for obj in tables)


def test_parse_pdf(tmp_path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    sample = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(sample), pagesize=letter)
    c.drawString(72, 720, "1 Introduction")
    c.drawString(72, 700, "This is a simple PDF line.")
    c.showPage()
    c.drawString(72, 720, "2 Requirements")
    c.drawString(72, 700, "Provide two bolts per assembly.")
    c.save()

    objects = parse_document(sample)

    assert any(obj["type"] == "text" for obj in objects)
    contents = "\n".join(obj["content"] for obj in objects if obj["type"] == "text")
    assert "Introduction" in contents
    assert "Provide two bolts" in contents
